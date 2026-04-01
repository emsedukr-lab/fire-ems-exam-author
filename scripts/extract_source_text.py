#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import shutil
import subprocess
import uuid
import zipfile
from dataclasses import asdict, dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

import pytesseract
from docx import Document
from lxml import etree
from openpyxl import load_workbook
from PIL import Image, ImageOps
from pypdf import PdfReader


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg"}
TEXT_EXTENSIONS = {".md", ".txt"}
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".hwp", ".hwpx"} | IMAGE_EXTENSIONS | TEXT_EXTENSIONS


def utc_now() -> str:
    return datetime.now(timezone.utc).replace(microsecond=0).isoformat()


@dataclass
class ManifestRecord:
    id: str
    filename: str
    copied_path: str
    media_type: str
    extractor: str
    text_path: str | None
    status: str
    fallback_used: str | None
    notes: list[str]
    source_refs: list[dict[str, str]]


def ensure_workspace(root: Path) -> None:
    for relative in ["sources", "sources/extracted", "bank", "outputs", "review"]:
        (root / relative).mkdir(parents=True, exist_ok=True)

    manifest_path = root / "sources" / "intake-manifest.json"
    if not manifest_path.exists():
        manifest_path.write_text(
            json.dumps({"version": 1, "generated_at": utc_now(), "source_documents": []}, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )


def load_manifest(path: Path) -> dict:
    if not path.exists():
        return {"version": 1, "generated_at": utc_now(), "source_documents": []}
    return json.loads(path.read_text(encoding="utf-8"))


def save_manifest(path: Path, payload: dict) -> None:
    payload["generated_at"] = utc_now()
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def copy_source(src: Path, dest_dir: Path) -> Path:
    dest = dest_dir / src.name
    if dest.exists():
        stem = src.stem
        suffix = src.suffix
        counter = 1
        while dest.exists():
            dest = dest_dir / f"{stem}-{counter}{suffix}"
            counter += 1
    shutil.copy2(src, dest)
    return dest


def extract_pdf(path: Path) -> tuple[str | None, list[str], str]:
    notes: list[str] = []
    text_parts: list[str] = []
    reader = PdfReader(str(path))
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            text_parts.append(f"## Page {index}\n{page_text.strip()}\n")
    if text_parts:
        return "\n".join(text_parts).strip() + "\n", notes, "pypdf"
    notes.append("No digital PDF text extracted. OCR or PDF re-export may be required.")
    return None, notes, "pypdf"


def extract_image(path: Path) -> tuple[str | None, list[str], str]:
    image = ImageOps.exif_transpose(Image.open(path))
    raw_text = pytesseract.image_to_string(image, lang="kor+eng", config="--psm 6")

    preprocessed = ImageOps.autocontrast(image.convert("L"))
    if preprocessed.width < 1800:
        preprocessed = preprocessed.resize((preprocessed.width * 2, preprocessed.height * 2))
    processed_text = pytesseract.image_to_string(preprocessed, lang="kor+eng", config="--psm 6")

    text = processed_text if len(processed_text.strip()) >= len(raw_text.strip()) else raw_text
    notes: list[str] = []
    if not text.strip():
        notes.append("OCR returned empty text.")
        return None, notes, "pytesseract"
    if text is processed_text and processed_text.strip() != raw_text.strip():
        notes.append("Used preprocessed OCR pass for improved extraction.")
    return text.strip() + "\n", notes, "pytesseract"


def extract_docx(path: Path) -> tuple[str | None, list[str], str]:
    notes: list[str] = []
    document = Document(str(path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            lines.append(paragraph.text.strip())
    for table_index, table in enumerate(document.tables, start=1):
        lines.append(f"\n## Table {table_index}")
        for row in table.rows:
            values = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append("\t".join(values))
    if lines:
        return "\n".join(lines).strip() + "\n", notes, "python-docx"

    textutil = shutil.which("textutil")
    if textutil:
        result = subprocess.run(
            [textutil, "-convert", "txt", "-stdout", str(path)],
            capture_output=True,
            text=True,
            check=False,
        )
        if result.returncode == 0 and result.stdout.strip():
            notes.append("Structured DOCX text was empty; used textutil fallback.")
            return result.stdout.strip() + "\n", notes, "textutil"

    notes.append("DOCX extraction returned no text.")
    return None, notes, "python-docx"


def extract_xlsx(path: Path) -> tuple[str | None, list[str], str]:
    notes: list[str] = []
    workbook = load_workbook(filename=str(path), read_only=True, data_only=False)
    blocks: list[str] = []
    for sheet in workbook.worksheets:
        blocks.append(f"## Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            if row is None:
                continue
            values = ["" if cell is None else str(cell).strip() for cell in row]
            if any(values):
                blocks.append("\t".join(values))
        blocks.append("")
    if blocks:
        return "\n".join(blocks).strip() + "\n", notes, "openpyxl"

    notes.append("Workbook opened but no text-like rows were extracted.")
    return None, notes, "openpyxl"


def extract_plain_text(path: Path) -> tuple[str | None, list[str], str]:
    notes: list[str] = []
    for encoding in ("utf-8", "utf-8-sig", "cp949"):
        try:
            text = path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        normalized = text.replace("\r\n", "\n").strip()
        if not normalized:
            notes.append("Text file was readable but empty.")
            return None, notes, "plain-text"
        if path.suffix.lower() == ".md":
            notes.append("Markdown formatting was preserved as plain text.")
        return normalized + "\n", notes, "plain-text"

    notes.append("Text file could not be decoded with utf-8 or cp949.")
    return None, notes, "plain-text"


def extract_hwpx(path: Path) -> tuple[str | None, list[str], str]:
    notes: list[str] = []
    fragments: list[str] = []
    with zipfile.ZipFile(path) as archive:
        xml_names = [name for name in archive.namelist() if name.lower().endswith(".xml")]
        if not xml_names:
            notes.append("HWPX archive contained no XML entries.")
            return None, notes, "hwpx-zip-xml"

        for name in xml_names:
            try:
                xml_bytes = archive.read(name)
                root = etree.fromstring(xml_bytes)
            except Exception:
                continue
            text = " ".join(token.strip() for token in root.itertext() if token.strip())
            if text:
                fragments.append(f"## {name}\n{text}\n")
    if fragments:
        return "\n".join(fragments).strip() + "\n", notes, "hwpx-zip-xml"

    notes.append("HWPX XML parse succeeded but yielded no usable text.")
    return None, notes, "hwpx-zip-xml"


def extract_hwp(path: Path) -> tuple[str | None, list[str], str, str | None, str]:
    notes: list[str] = []
    for candidate in ["hwp5txt", "hwp5proc"]:
        executable = shutil.which(candidate)
        if not executable:
            continue
        result = subprocess.run([executable, str(path)], capture_output=True, text=True, check=False)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip() + "\n", notes, candidate, None, "extracted"
        notes.append(f"{candidate} was available but did not return usable text.")

    notes.append("No dedicated HWP parser succeeded. Request PDF or HWPX re-export.")
    return None, notes, "hwp-conversion-required", "pdf_reexport", "conversion_required"


def extract_with_route(path: Path) -> tuple[str | None, list[str], str, str | None, str]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        text, notes, extractor = extract_pdf(path)
        return text, notes, extractor, None, "extracted" if text else "partial"
    if suffix in IMAGE_EXTENSIONS:
        text, notes, extractor = extract_image(path)
        return text, notes, extractor, None, "extracted" if text else "partial"
    if suffix == ".docx":
        text, notes, extractor = extract_docx(path)
        return text, notes, extractor, None, "extracted" if text else "partial"
    if suffix == ".xlsx":
        text, notes, extractor = extract_xlsx(path)
        return text, notes, extractor, None, "extracted" if text else "partial"
    if suffix in TEXT_EXTENSIONS:
        text, notes, extractor = extract_plain_text(path)
        return text, notes, extractor, None, "extracted" if text else "partial"
    if suffix == ".hwpx":
        text, notes, extractor = extract_hwpx(path)
        status = "extracted" if text else "partial"
        fallback = None if text else "pdf_or_docx_export"
        return text, notes, extractor, fallback, status
    if suffix == ".hwp":
        return extract_hwp(path)
    return None, [f"Unsupported source type: {suffix or '(no extension)'}"], "unsupported", None, "failed"


def write_text_output(path: Path, text: str) -> None:
    path.write_text(text, encoding="utf-8")


def update_manifest(manifest_path: Path, records: Iterable[ManifestRecord]) -> None:
    manifest = load_manifest(manifest_path)
    existing = manifest.get("source_documents", [])
    existing.extend(asdict(record) for record in records)
    manifest["source_documents"] = existing
    save_manifest(manifest_path, manifest)


def main() -> int:
    parser = argparse.ArgumentParser(description="Copy exam source files into the workspace and extract text when possible.")
    parser.add_argument("sources", nargs="+", help="Source files to ingest.")
    parser.add_argument("--workspace", default=".", help="Workspace root. Defaults to current directory.")
    args = parser.parse_args()

    workspace = Path(args.workspace).expanduser().resolve()
    ensure_workspace(workspace)

    manifest_path = workspace / "sources" / "intake-manifest.json"
    copied_dir = workspace / "sources"
    extracted_dir = workspace / "sources" / "extracted"

    records: list[ManifestRecord] = []
    for raw_source in args.sources:
        source_path = Path(raw_source).expanduser().resolve()
        if not source_path.exists():
            raise FileNotFoundError(f"Source file not found: {source_path}")

        record_id = f"src-{uuid.uuid4().hex[:8]}"
        copied_path = copy_source(source_path, copied_dir)
        text, notes, extractor, fallback_used, status = extract_with_route(copied_path)

        text_output = None
        if text:
            text_output = extracted_dir / f"{record_id}.txt"
            write_text_output(text_output, text)

        records.append(
            ManifestRecord(
                id=record_id,
                filename=source_path.name,
                copied_path=str(copied_path.relative_to(workspace)),
                media_type=source_path.suffix.lower().lstrip("."),
                extractor=extractor,
                text_path=str(text_output.relative_to(workspace)) if text_output else None,
                status=status,
                fallback_used=fallback_used,
                notes=notes,
                source_refs=[{"kind": "file", "value": str(copied_path.relative_to(workspace))}],
            )
        )

    update_manifest(manifest_path, records)
    print(json.dumps([asdict(record) for record in records], ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
